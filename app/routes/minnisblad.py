"""This module contains the routes for the minnisblad application."""

from datetime import datetime
import json
from tempfile import NamedTemporaryFile
from fastapi import (
    APIRouter,
    Request,
    File,
    UploadFile,
    Form,
    Depends,
    HTTPException,
)
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.templating import Jinja2Templates
from docx import Document
from app.utils import (
    send_text_to_openai,
    get_token,
    process_uploaded_file,
)


templates = Jinja2Templates(directory="app/templates")

router = APIRouter()


@router.get("/minnisblad", response_class=HTMLResponse)
async def minnisblad(request: Request):
    """Render the minnisblad page."""
    return templates.TemplateResponse("minnisblad.html", {"request": request})


@router.post("/minnisblad/upload/")
async def upload_file(
    file: UploadFile = File(...),
    chapters: str = Form(...),
    _: str = Depends(get_token),
):
    """Process the uploaded file and return the modified document."""
    text = await process_uploaded_file(file)
    try:
        selected_chapters = json.loads(chapters)
        respond_format = create_response_format(selected_chapters)
        openai_response = await send_text_to_openai(text, respond_format)
        output_file_path = create_docx_from_json(openai_response)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = "Frodi_minnisblad_" + timestamp + ".docx"

        return FileResponse(
            output_file_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename,
        )
    except Exception as e:
        raise HTTPException(
            status_code=500, detail="An unexpected error occurred"
        ) from e


def create_response_format(selected_chapters: list) -> dict:
    """Create the response format based on the selected chapters."""
    base_response = {
        "type": "json_schema",
        "json_schema": {
            "name": "minnisblad",
            "strict": True,
            "schema": {
                "type": "object",
                "properties": {
                    "titill": {
                        "type": "string",
                        "description": "Titill minnisblaðsins.",
                    },
                    "kaflar": {
                        "type": "array",
                        "description": "Kaflar minnisblaðsins.",
                        "items": {
                            "type": "object",
                            "properties": {
                                "chapter_title": {
                                    "type": "string",
                                    "description": "Titill kafla.",
                                },
                                "content": {
                                    "type": "string",
                                    "description": "Innihald kafla.",
                                },
                            },
                            "required": ["chapter_title", "content"],
                            "additionalProperties": False,
                        },
                    },
                },
                "required": [
                    "titill",
                    "kaflar",
                ],
                "additionalProperties": False,
            },
        },
    }
    chapter_descriptions = {
        "inngangur": "Inngangur minnisblaðsins.",
        "samantekt": "Samantekt minnisblaðsins.",
        "aaetlun": "Áætlun minnisblaðsins.",
        "markmid": "Markmið minnisblaðsins.",
    }

    for chapter in selected_chapters:
        if chapter in chapter_descriptions:
            base_response["json_schema"]["schema"]["properties"][chapter] = {
                "type": "string",
                "description": chapter_descriptions[chapter],
            }
            base_response["json_schema"]["schema"]["required"].append(chapter)

    return base_response


def create_docx_from_json(response_json: dict) -> str:
    """Create a Word document from the JSON response."""
    doc = Document()

    doc.add_heading(response_json.get("titill", "Titill Ekki Tiltækur"), level=1)

    if "Inngangur" in response_json:
        doc.add_heading("Inngangur", level=2)
        doc.add_paragraph(response_json["Inngangur"])

    if "markmid" in response_json:
        doc.add_heading("Markmið", level=2)
        doc.add_paragraph(response_json["markmid"])

    if "aaetlun" in response_json:
        doc.add_heading("Áætlun", level=2)
        doc.add_paragraph(response_json["aaetlun"])

    if "kaflar" in response_json:
        for chapter in response_json["kaflar"]:
            doc.add_heading(chapter["chapter_title"], level=2)
            doc.add_paragraph(chapter["content"])

    if "Samantekt" in response_json:
        doc.add_heading("Samantekt", level=2)
        doc.add_paragraph(response_json["Samantekt"])

    with NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        doc.save(temp_file.name)
        temp_file_path = temp_file.name

    return temp_file_path
